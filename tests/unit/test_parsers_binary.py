"""Parser tests for binary formats (PDF, DOCX) that need a real generated
file rather than hand-written bytes."""

import io

import docx
import pymupdf

from agentic_rag.ingestion.parsed_document import ElementType
from agentic_rag.ingestion.parsers.docx import DocxParser
from agentic_rag.ingestion.parsers.pdf import PdfParser


def _build_sample_pdf() -> bytes:
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Document Heading", fontsize=20)
    page.insert_text((72, 110), "This is a normal paragraph of body text.", fontsize=10)
    data = doc.tobytes()
    doc.close()
    return data


def _build_pdf_with_bolded_word_inside_a_paragraph() -> bytes:
    """One block: a long paragraph whose first word happens to render in a
    much larger font (e.g. a defined term) — should stay one paragraph,
    not become a "heading" just because of that one word."""
    doc = pymupdf.open()
    page = doc.new_page()
    html = (
        '<p style="font-size:10pt">'
        '<span style="font-size:20pt">Landlord</span>'
        " covenants and agrees that this entire sentence is ordinary body"
        " text describing a real obligation under the lease, long enough"
        " that it could never be mistaken for a section heading on its own,"
        " and it must not be reclassified as one just because a single"
        " defined term at the start happens to render in a larger font."
        "</p>"
    )
    page.insert_htmlbox(pymupdf.Rect(72, 72, 500, 300), html)
    data = doc.tobytes()
    doc.close()
    return data


def _build_pdf_with_whitespace_only_large_span() -> bytes:
    """One block: ordinary body text with a blank/whitespace-only run at a
    much larger font size in the middle — a real artifact found in a real
    PDF, where it silently inflated the block's max font size without
    contributing any actual visible text."""
    doc = pymupdf.open()
    page = doc.new_page()
    html = (
        '<p style="font-size:10pt">Ordinary paragraph text continues'
        '<span style="font-size:30pt"> </span>'
        "across this blank run with no visible effect on the reader"
        " and should stay a normal paragraph.</p>"
    )
    page.insert_htmlbox(pymupdf.Rect(72, 72, 500, 300), html)
    data = doc.tobytes()
    doc.close()
    return data


def test_pdf_parser_does_not_reclassify_a_paragraph_with_one_bolded_word() -> None:
    parsed = PdfParser().parse(
        filename="a.pdf", content=_build_pdf_with_bolded_word_inside_a_paragraph()
    )
    assert len(parsed.elements) == 1
    assert parsed.elements[0].element_type == ElementType.PARAGRAPH
    assert "covenants and agrees" in parsed.elements[0].text


def test_pdf_parser_ignores_whitespace_only_spans_for_heading_detection() -> None:
    parsed = PdfParser().parse(
        filename="a.pdf", content=_build_pdf_with_whitespace_only_large_span()
    )
    assert parsed.elements
    assert all(e.element_type == ElementType.PARAGRAPH for e in parsed.elements)


def test_pdf_parser_distinguishes_heading_by_font_size() -> None:
    parsed = PdfParser().parse(filename="a.pdf", content=_build_sample_pdf())
    assert parsed.page_count == 1

    headings = [e for e in parsed.elements if e.element_type == ElementType.HEADING]
    paragraphs = [e for e in parsed.elements if e.element_type == ElementType.PARAGRAPH]

    assert any("Document Heading" in h.text for h in headings)
    assert any("normal paragraph" in p.text for p in paragraphs)
    assert all(e.page == 1 for e in parsed.elements)


def _build_sample_docx() -> bytes:
    document = docx.Document()
    document.add_heading("Report Title", level=1)
    document.add_paragraph("Body paragraph under the heading.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Key"
    table.rows[0].cells[1].text = "Value"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_parser_extracts_headings_paragraphs_and_tables() -> None:
    parsed = DocxParser().parse(filename="a.docx", content=_build_sample_docx())

    headings = [e for e in parsed.elements if e.element_type == ElementType.HEADING]
    paragraphs = [e for e in parsed.elements if e.element_type == ElementType.PARAGRAPH]
    tables = [e for e in parsed.elements if e.element_type == ElementType.TABLE]

    assert headings[0].text == "Report Title"
    assert paragraphs[0].text == "Body paragraph under the heading."
    assert paragraphs[0].heading == "Report Title"
    assert "Key | Value" in tables[0].text
