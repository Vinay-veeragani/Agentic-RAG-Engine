from __future__ import annotations

import io

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from agentic_rag.core.errors import InvalidDocumentError
from agentic_rag.ingestion.cleaners.text import clean_text
from agentic_rag.ingestion.parsed_document import DocumentElement, ElementType, ParsedDocument


class DocxParser:
    """Walks the document body's paragraphs and tables in document order
    (python-docx's `.paragraphs`/`.tables` accessors lose relative order,
    so this iterates the underlying XML body directly instead)."""

    def parse(self, *, filename: str, content: bytes) -> ParsedDocument:
        try:
            document = docx.Document(io.BytesIO(content))
        except Exception as exc:  # python-docx raises assorted low-level errors
            raise InvalidDocumentError(
                f"could not parse {filename!r} as a .docx file", details={"filename": filename}
            ) from exc

        elements: list[DocumentElement] = []
        current_heading: str | None = None
        order_index = 0

        for child in document.element.body.iterchildren():
            if child.tag == qn("w:p"):
                paragraph = Paragraph(child, document)
                text = clean_text(paragraph.text)
                if not text:
                    continue
                style_name = (paragraph.style.name or "") if paragraph.style else ""
                if style_name.lower().startswith("heading"):
                    current_heading = text
                    elements.append(
                        DocumentElement(
                            element_type=ElementType.HEADING,
                            text=text,
                            order_index=order_index,
                            heading=current_heading,
                            metadata={"style": style_name},
                        )
                    )
                else:
                    elements.append(
                        DocumentElement(
                            element_type=ElementType.PARAGRAPH,
                            text=text,
                            order_index=order_index,
                            heading=current_heading,
                        )
                    )
                order_index += 1

            elif child.tag == qn("w:tbl"):
                table = Table(child, document)
                table_text = _render_table(table)
                if table_text:
                    elements.append(
                        DocumentElement(
                            element_type=ElementType.TABLE,
                            text=table_text,
                            order_index=order_index,
                            heading=current_heading,
                        )
                    )
                    order_index += 1

        title = document.core_properties.title or None
        return ParsedDocument(
            filename=filename, document_type="docx", elements=elements, title=title
        )


def _render_table(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [clean_text(cell.text) for cell in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)
